import os
from docx import Document

class Sentence:
    def __init__(self, protocol_name, keneset, protocol_type, protocol_no, speaker, text):
        self.protocol_name = protocol_name
        self.keneset = keneset
        self.protocol_type = protocol_type
        self.protocol_no = protocol_no
        self.speaker = speaker
        self.text = text
    
class Protocol:
    def __init__(self, name, keneset, type, number):
        self.name = name
        self.keneset = keneset
        self.type = type
        self.number = number
        self.sentences = {}
    def add_sentence(self, speaker, sentence):
        if speaker in self.sentences:
            self.sentences[speaker].append(sentence)
        else:
            self.sentences[speaker] = [sentence]
    def print_speakers(self):
        print(f"{self.name}:")
        for speaker in self.sentences.keys():
            print(speaker)
    def print_speech(self):
        print("Speech:")
        for sentences in self.sentences.values():
            for sentence in sentences:
                print(sentence.text)

#Input: String of number which can be numeric or as hebrew word
#Output: Integer equivalent of the number
def convertToInt(word):
    digits_dict = {'אחת': 1, 'שתיים': 2, 'שלוש': 3, 'ארבע': 4, 'חמש': 5, 'חמיש': 5, 'שש': 6, 'שיש': 6, 'שבע': 7, 'שמונה': 8, 'תשע': 9, 'עשר': 10, 'עשרים': 20, 'שמונים': 80, 'מאה': 100, 'מאתיים': 200}
    if word[:-1].isdigit():
        if word[-1].isdigit():
            return int(word)
        else:
            return int(word[:-1]) #to get rid of '<' at the end of the number
    else:
        word = word.replace('-', ' ') #consider '-' as space since noticed that there are many protocols numbers written as strings consists of words seperated by '-' 
        word_splitted = word.split()
        word_splitted = [word[1:] if word[0] in ['ו', 'ה']else word for word in word_splitted]
        for i in range(len(word_splitted)):
            if word_splitted[i] in digits_dict:
                word_splitted[i] = digits_dict[word_splitted[i]]
            elif word_splitted[i][-2:] == 'ים' and word_splitted[i][:-2] in digits_dict:
                word_splitted[i] = digits_dict[word_splitted[i][:-2]] * 10
            elif word_splitted[i] == 'מאות' and i > 0:
                word_splitted[i-1] *= 100
                word_splitted[i] = 0
            elif word_splitted[i] == 'עשרה' and i > 0:
                word_splitted[i-1] += 10
                word_splitted[i] = 0
            else:
                return -1
        return sum(word_splitted)

#Input: string which could represent someone's name 
#Output: the name without the additions and titles 
def speakerClean(speaker):
    ministries = ['התשתיות הלאומיות', 'המשטרה', 'לביטחון פנים', 'לאיכות הסביבה', 'להגנת הסביבה', 'החינוך והתרבות', 'החינוך', 'התחבורה והבטיחות בדרכים', 'התחבורה', 'האוצר', 'הכלכלה והתכנון', 'המשפטים', 'הבריאות', 'החקלאות ופיתוח הכפר', 'החקלאות', 'החוץ', 'הבינוי והשיכון ', 'העבודה, הרווחה והשירותים החברתיים', 'העבודה', 'התעשייה והמסחר', 'התעשייה, המסחר והתעסוקה', 'התיירות', 'המדע והטכנולוגיה', 'הפנים', 'המדע, התרבות והספורט', 'התרבות והספורט', 'האנרגיה והמים', 'לענייני דתות', 'במשרד ראש הממשלה', 'לנושאים אסטרטגיים ולענייני מודיעין', 'לקליטת העלייה', 'לאזרחים ותיקים']
    open_brac = speaker.find('(')
    close_brac = speaker.find(')')
    if open_brac != -1 and close_brac != -1:
        speaker = speaker[:open_brac] + speaker[close_brac+1:]
    open_an_brac = speaker.find('<<')
    close_an_brac = speaker.find('>>')
    while(open_an_brac != -1 and close_an_brac != -1):
        speaker = speaker[:open_an_brac] + speaker[close_an_brac+1:]
        open_an_brac = speaker.find('<<')
        close_an_brac = speaker.find('>>')
    speaker = speaker.replace('<', '').replace('>','').replace('היו"ר', '').replace('היו”ר', '').replace('יו"ר הכנסת', '').replace('יו”ר הכנסת', '').replace('יו"ר ועדת הכנסת', '').replace('יו”ר ועדת הכנסת', '').replace('-', '').replace('מ"מ', '').replace('מ”מ', '').replace('סגן', '').replace('סגנית', '').replace('מזכיר הכנסת','').replace('מזכירת הכנסת','').replace('תשובת','').replace('ראש הממשלה', '')
    if 'שר' in speaker or 'השר' in speaker or 'שרת' in speaker or 'השרה' in speaker:     
        speaker = speaker.replace('שר ','').replace('השר ','').replace('שרת ','').replace('השרה ','')
        for ministry in ministries:
            if ministry in speaker or 'ה' + ministry in speaker:
                speaker = speaker.replace(ministry, '')
    speaker = speaker.strip()
    return speaker

#Input: file/protocol name
# Output: The keneset number to where the protocol relates, protocol type       
def extract_metada_from_name(file_name):
    file_name_splitted = file_name.split('_')
    keneset_no = int(file_name_splitted[0])
    protocol_type = 'committee' if file_name_splitted[1][-1] == 'v' else 'plenary' if file_name_splitted[1][-1] == 'm' else None
    return keneset_no, protocol_type

#Input: .docx file of a protocol
#Output: protocol number if found, -1 otherwise. 
def extract_metada_from_content(file_content):
    for paragraph in file_content.paragraphs:
        paragraph_stripped = paragraph.text.strip()
        for word in ["הישיבה", "פרוטוקול מס'"]:
            while True:
                word_idx = paragraph_stripped.find(word)
                if word_idx != -1:
                    paragraph_stripped = paragraph_stripped[word_idx+len(word):]
                    paragraph_words = paragraph_stripped.split()
                    if paragraph_words:
                        string_to_int = convertToInt(paragraph_words[0])
                    else:
                        continue
                    if string_to_int != -1: 
                        return string_to_int 
                    else:
                        continue 
                else:
                    break
    return -1

#helper method for for extract_relevant_text
#Input: .docx file of a protocol
#Output: index of the first relevant paragraph
def find_starting_relevant(file_content): 
    paragraph_idx = 0
    for paragraph_idx in range(len(file_content.paragraphs)):
        paragraph_txt = file_content.paragraphs[paragraph_idx].text.strip()
        if (paragraph_txt.startswith('היו"ר') or paragraph_txt.startswith('היו”ר') or paragraph_txt.startswith('יו"ר הכנסת') or paragraph_txt.startswith('יו”ר הכנסת') or paragraph_txt.startswith('מ"מ היו"ר') or paragraph_txt.startswith('מ”מ היו”ר')) and paragraph_txt.endswith(':'):
            return paragraph_idx
        elif paragraph_txt.startswith('<< יור >>') and paragraph_txt.endswith('<< יור >>'):
            return paragraph_idx
        elif paragraph_txt.startswith('<היו"ר') and paragraph_txt.endswith(':>'):
            return paragraph_idx
    return paragraph_idx

#helper method for for extract_relevant_text
#Input: .docx file of a protocol
#Output: index of the last relevant paragraph
def find_last_relevant(file_content): 
    idx = len(file_content.paragraphs) - 1
    curr_paragraph = file_content.paragraphs[idx].text
    while (not len(curr_paragraph) or ('הישיבה ננעלה' not in curr_paragraph and 'הטקס ננעל' not in curr_paragraph)) and idx > 0:
        idx -= 1
        curr_paragraph = file_content.paragraphs[idx].text
    if idx == 0:
        return idx - 1 #when failed to find last relevant, the last sentence will be the last relevant
    else: 
        return idx

#helper method for sentence_handle which checks validity of a sentence
def sentence_validity(sentence):
    hebrew_alphabet = {'א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י', 'כ', 'ל', 'מ', 'נ', 'ס', 'ע', 'פ', 'צ', 'ק', 'ר', 'ש', 'ת', '', 'ך', 'ם', 'ן', 'ף', 'ץ'}
    if all(letter not in hebrew_alphabet for letter in sentence):
        return False
    if any('a' <= letter <= 'z' or 'A' <= letter <= 'Z' for letter in sentence):
        return False
    if '- - -' in sentence or '---' in sentence or '– – –' in sentence or '–––' in sentence:
        return False
    return True
    

#helper method for for extract_relevant_text that creates a sentence object and add it to the relevant protocol 
def sentence_handle(protocol, curr_speaker, paragraph_txt):
    curr_sentence = ''
    for letter in paragraph_txt: 
        if letter in ['.', '?', '!']:
            curr_sentence = curr_sentence.strip()
            if sentence_validity(curr_sentence):
                sentence = Sentence(protocol.name, protocol.keneset, protocol.type, protocol.number, curr_speaker, curr_sentence)
                protocol.add_sentence(curr_speaker, sentence)
            curr_sentence = ''
        else:
            curr_sentence += letter

#extracts the relevant sentences and arranges them according to protocol and speaker
def extract_relevant_text(file_content, protocol):
    first_idx = find_starting_relevant(file_content)
    last_idx = find_last_relevant(file_content)
    curr_speaker = None
    for paragraph in file_content.paragraphs[first_idx:last_idx]:
        paragraph_txt = paragraph.text.strip()
        if paragraph_txt.endswith(':'):
            paragraph_txt_cleaned = speakerClean(paragraph_txt)[:-1]
            if len(paragraph_txt_cleaned.split()) < 6:
                curr_speaker = paragraph_txt_cleaned
            else: #deal with it as speech
                sentence_handle(protocol, curr_speaker, paragraph_txt)
        elif paragraph_txt.endswith(':>'):
            paragraph_txt_cleaned = speakerClean(paragraph_txt)[:-1]
            if len(paragraph_txt_cleaned.split()) < 6:
                curr_speaker = paragraph_txt_cleaned
            else: #deal with it as speech
                sentence_handle(protocol, curr_speaker, paragraph_txt)
        elif ':' in paragraph_txt:
            pot_curr_speaker = speakerClean(paragraph_txt)
            if pot_curr_speaker.endswith(':'):
                paragraph_txt_cleaned = pot_curr_speaker[:-1]
                if len(paragraph_txt_cleaned.split()) < 6:
                    curr_speaker = paragraph_txt_cleaned
                else: #deal with it as speech
                    sentence_handle(protocol, curr_speaker, paragraph_txt)
            else: #else deal with it as a speech 
                sentence_handle(protocol, curr_speaker, paragraph_txt)
        else:
            sentence_handle(protocol, curr_speaker, paragraph_txt)
    #don't forget the case of the file where it will return nothing 

#Input: path to a folder
#Output: file names, paths, and contents for all .docx files in the folder
def read_files(folder):
    file_names = [file for file in os.listdir(folder) if file[-5:] == '.docx']
    file_paths = {filename: os.path.join(folder, filename) for filename in file_names}
    file_contents = {file_name: Document(file_path) for file_name, file_path in file_paths.items()}
    return file_names, file_paths, file_contents


folder_path = "protocol_for_hw1" 
file_names, file_paths, file_contents = read_files(folder_path)
protocols = []

for file_name in file_names: 
    keneset_no, protocol_type = extract_metada_from_name(file_name)
    protocol_no = extract_metada_from_content(file_contents[file_name])
    protocol = Protocol(file_name, keneset_no, protocol_type, protocol_no)
    protocols.append(protocol)
    extract_relevant_text(file_contents[file_name], protocol)
    print(file_name)

    #protocol.print_speakers()
    protocol.print_speech()

