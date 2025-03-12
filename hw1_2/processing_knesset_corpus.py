import os
from docx import Document
import json
import sys
import re 

class Sentence:
    def __init__(self, protocol_name, keneset, protocol_type, protocol_no, speaker, text):
        self.protocol_name = protocol_name
        self.keneset = keneset
        self.protocol_type = protocol_type
        self.protocol_no = protocol_no
        self.speaker = speaker
        self.text = text
    
class Protocol:
    def __init__(self, name, keneset, type, number):
        self.name = name
        self.keneset = keneset
        self.type = type
        self.number = number
        self.sentences = {}
    def add_sentence(self, speaker, sentence):
        if speaker in self.sentences:
            self.sentences[speaker].append(sentence)
        else:
            self.sentences[speaker] = [sentence]

ministries = {'התשתיות', 'הלאומיות', 'האנרגיה', 'המים', 'המשטרה', 'ביטחון', 'פנים', 'איכות', 'הגנת', 
              'הסביבה', 'החינוך', 'התרבות', 'הספורט', 'המדע', 'הטכנולוגיה', 'התחבורה', 'הבטיחות', 
              'בדרכים', 'האוצר', 'הכלכלה', 'התכנון', 'התעשייה', 'המסחר', 'התעסוקה', 'התיירות', 
              'הבריאות', 'החקלאות', 'פיתוח', 'הכפר', 'הבינוי', 'השיכון', 'העבודה', 'הרווחה', 
              'השירותים', 'החברתיים', 'לענייני', 'דתות', 'במשרד', 'ראש', 'הממשלה', 'לנושאים', 
              'אסטרטגיים', 'מודיעין', 'לקליטת', 'עלייה', 'לאזרחים', 'ותיקים', 'התקשורת', 'משפטים', 
              'חוץ', 'לאומי', 'שוויון', 'חברתי', 'מעמד', 'האישה', 'קידום', 'שירותי', 'דת', 'שיתוף', 
              'פעולה', 'אזורי', 'ההתיישבות', 'המשימות', 'המזון', 'החדשנות', 'הנגב', 'הגליל', 'החוסן', 
              'הלאומי', 'הקליטה', 'תפוצות', 'התפוצות', 'המאבק', 'באנטישמיות', 'מורשת', 'חיזוק', 'קידום', 
              'קהילתי', 'הפריפריה', 'החלופי', 'הדיגיטל', 'ההשכלה', 'הגבוהה', 'המשלימה', 'משאבי', 'החלל', 
              'רה"מ', 'ירושלים', 'המקשר', 'בין', 'הממשלה', 'לכנסת', 'העורף', 'ההסברה', 'מיעטום', 'חברה', 
              'גמלאים', 'האמנויות'}
titles_and_symbols = ['<', '>', 'היו"ר', 'היו”ר', 'יו"ר הכנסת', 'יו”ר הכנסת', 
                      'יו"ר ועדת הכנסת', 'יו”ר ועדת הכנסת', 'מ"מ', 'מ”מ', 'סגן', 'סגנית', 
                      'מזכיר הכנסת', 'מזכירת הכנסת', 'תשובת', 'המשנה לראש הממשלה', 'ראש הממשלה', 
                      'עו"ד', 'עו”ד', 'ד"ר', 'ד”ר', "פרופ'", 'נצ"מ', 'ניצב']
invalid_names = {'ביום', 'קריאות', 'קריאה'}
digits_dict = {'אחת': 1, 'שתיים': 2, 'שתים' : 2, 'שלוש': 3, 'ארבע': 4, 'חמש': 5, 'חמיש': 5, 
               'שש': 6, 'שיש': 6, 'שבע': 7, 'שמונה': 8, 'תשע': 9, 'עשר': 10, 'עשרים': 20, 
               'שמונים': 80, 'מאה': 100, 'מאתיים': 200, 'אלף': 1000}

#Helper method for 'extract_metada_from_content'. It checks if the given string represent an integer number, and if yes, it return its integer value
#Input: String of number which can be numeric or as hebrew word
#Output: Integer equivalent of the number or -1 in case it's not identified as a number
def convertToInt(word):
    # case 1: number in numeric form
    if word.isdigit():
        return int(word)
    elif word[:-1].isdigit(): #to deal with the cases where the number followed by '<' directly
        return int(word[:-1])
    #case 2: number in string form 
    else:
        word_splitted = word.split('-')
        word_splitted = [word[1:] if word[0] in ['ו', 'ה']else word for word in word_splitted]
        for i in range(len(word_splitted)):
            #option 1: digit from the digits dictionary 
            if word_splitted[i] in digits_dict:
                word_splitted[i] = digits_dict[word_splitted[i]]
            #option 2: digit in plural form
            elif word_splitted[i][-2:] == 'ים' and word_splitted[i][:-2] in digits_dict:
                word_splitted[i] = digits_dict[word_splitted[i][:-2]] * 10
            #option 3: the word 'מאות'
            elif word_splitted[i] == 'מאות' and i > 0:
                word_splitted[i-1] *= 100
                word_splitted[i] = 0
            #option 4: the word 'עשרה'
            elif word_splitted[i] == 'עשרה' and i > 0:
                word_splitted[i-1] += 10
                word_splitted[i] = 0
            #option 5: the word didn't identify as a number
            else:
                return -1
        return sum(word_splitted)

#Input: string which could represent someone's name 
#Output: the pure name (without titles, tags, and party name)
def speakerClean(speaker):
    #remove the party name
    speaker = re.sub(r"\(.*?\)", "", speaker)
    #remove the tag
    speaker = re.sub(r"<<.*?>>", "", speaker)
    #remove the title and other symbols
    for element in titles_and_symbols:
        speaker = speaker.replace(element, '')
    #replace the ',' by a space
    speaker = speaker.replace(',', ' ')
    #replace multi spaces by one space
    speaker = ' '.join(speaker.split())
    #strip and remove '-' at the beginning and end of a speaker's name 
    speaker = speaker.strip('-')
    #remove minister titles, these titles are special case because they may be followed by the ministry name 
    if speaker.find('שר ') == 0 or speaker.find('השר ') == 0 or speaker.find('השרה ') == 0 or speaker.find('שרת ') == 0:
        speaker = speaker.replace('השרה ', '').replace('שרת ', '').replace('השר ', '').replace('שר ', '')
        speaker_splitted = speaker.split()
        idx_to_cut = 0
        while idx_to_cut < len(speaker_splitted) and speaker_splitted[idx_to_cut] in ministries or (speaker_splitted[idx_to_cut][0] in ['ל', 'ו', 'ה'] and speaker_splitted[idx_to_cut][1:] in ministries):
            idx_to_cut += 1 
        speaker = ' '.join(speaker_splitted[idx_to_cut:])
    return speaker

#Input: file/protocol name
# Output: The keneset number to where the protocol relates, protocol type       
def extract_metada_from_name(file_name):
    try:
        file_name_splitted = file_name.split('_')
        keneset_no = int(file_name_splitted[0])
        protocol_type = 'committee' if file_name_splitted[1][-1] == 'v' else 'plenary' if file_name_splitted[1][-1] == 'm' else None
        return keneset_no, protocol_type
    except (IndexError, ValueError) as e:
        print(f"Failed to extract metadata from {file_name}: {e}") 

#Input: .docx file of a protocol
#Output: protocol number if found, -1 otherwise. 
def extract_metada_from_content(file_content):
    try:
        for paragraph in file_content.paragraphs:
            paragraph_stripped = paragraph.text.strip()
            #iterates over the occurrences of the two string until a one followed by a number found
            for word in ["הישיבה", "פרוטוקול מס'"]:
                while True:
                    try:
                        word_idx = paragraph_stripped.find(word)
                        if word_idx != -1:
                            paragraph_stripped = paragraph_stripped[word_idx+len(word):] #slicing the paragraph to the part after the occurence
                            paragraph_words = paragraph_stripped.split()
                            if paragraph_words:
                                string_to_int = convertToInt(paragraph_words[0]) #if the string followed by a number it should be the first in the new sliced paragraph
                            else:
                                continue
                            if string_to_int != -1: #if a number identified after the string
                                return string_to_int 
                            else:
                                continue 
                        else:
                            break
                    except Exception as e:
                        print(f"Error in procssing the paragraph text: {e}")
    except AttributeError as e:
        print(f"Invalid file content: {e}")
    except Exception as e:
        print(f"Unexpected error")
    return -1

#checks if the paragraph or part of it is underlined
def is_underlined(paragraph):
    #checking paragraph style for underline
    par_style = paragraph.style
    while par_style:
        if hasattr(par_style, 'font') and par_style.font.underline:
            return True
        par_style = par_style.base_style
    #Checking individual runs for underline
    for run in paragraph.runs:
        if run.font and run.font.underline:
            return True
    return False

#helper method for for extract_relevant_text, it identifies the start of the relevant text
#Input: .docx file of a protocol
#Output: index of the first relevant paragraph
def find_starting_relevant(document):
    skip_keywords = {
    'רישום פרלמנטרי', 'סדר היום', 'מוזמנים', 'משרד האוצר', 'נכחו',
    'רשמת פרלמנטרית', 'רשמה וערכה', 'ייעוץ משפטי', 'חברי הוועדה',
    'חברי כנסת', 'יועצת משפטית', 'יועץ משפטי', 'מנהל הוועדה',
    'מנהלת הוועדה', 'מזכירת הוועדה', 'הגילויים החדשים', 'הצבעה',
    'קצרנית', 'מנהל/ת הוועדה', 'משרד המשפטים', 'סדר-היום', 'נוכחים',
    'משתתפים (באמצעים מקוונים)', 'חברי הכנסת', 'משתתפים באמצעים מקוונים', 'מנהלות הוועדה', ':סדר היום', 'סדר היום:',
    'שינויים', 'הכנסת', 'הכנסת:', 'הכנסה', 'ועדה לדיון', 'הצעת חוק', 'מנהלי הוועדה', 'מנהלי הוועדה:'}
    try:
        #iterating over the paragraph until the new speaker is found 
        for idx, paragraph in enumerate(document.paragraphs):
            text = ' '.join(paragraph.text.split()).strip() #getting rid of multi-spaces 
            text = re.sub(r"<<[^<>]*?>>", '', text).strip() #getting rid of tags 
            if any(keyword in text for keyword in skip_keywords): #skipping common non-speakers underlined words 
                continue

            if (text.endswith(':') or text.endswith('>')) and is_underlined(paragraph):
                return idx

        return 0 
    except Exception as e:
        print(f"Error in finding the first relevant sentence: {e}")
        return 0


#helper method for for extract_relevant_text, it identifies the last relevant paragraph
#Input: .docx file of a protocol
#Output: index of the last relevant paragraph
def find_last_relevant(file_content): 
    idx = len(file_content.paragraphs) - 1
    curr_paragraph = file_content.paragraphs[idx].text
    while (not len(curr_paragraph) or ('הישיבה ננעלה' not in curr_paragraph and 'הטקס ננעל' not in curr_paragraph)) and idx > 0:
        idx -= 1
        curr_paragraph = file_content.paragraphs[idx].text
    if idx == 0: #when it fails to find a message states debate ending explicitly, it assumes that the last paragraph in the protocol is the last relevant
        return idx - 1 
    else: 
        return idx

#helper method for sentence_handle which checks validity of a sentence
def sentence_validity(sentence):
    if all(not('א' <= letter <= 'ת') for letter in sentence):
        return False
    if any('a' <= letter <= 'z' or 'A' <= letter <= 'Z' for letter in sentence):
        return False
    if '- -' in sentence or '– –' in sentence or '...' in sentence:
        return False
    return True

def sentence_tokenize(sentence):
    marks = {';', '(', ')', ' ', '-', '–', '!', '?', '.', '"', ',', ':'}
    tokenized_sentence = []
    word = ''
    for i,letter in enumerate(sentence):
        #the case when " is part of a word
        if letter == '"' and i != 0 and i != len(sentence) - 1 and 'א' <= sentence[i-1] <= 'ת' and 'א' <= sentence[i+1] <= 'ת':
            word += letter
        #thse case when , is part of a number
        elif letter == ',' and i != 0 and i != len(sentence) - 1 and sentence[i-1].isdigit() and sentence[i+1].isdigit():
            word += letter
        #the case when : is part of a term indicating time
        elif letter == ':' and i != 0 and i <= len(sentence) - 3 and sentence[i-1].isdigit() and sentence[i+1].isdigit() and sentence[i+2].isdigit():
            word += letter 
        #the case when . is decimal point or part of date term
        elif letter == '.' and i != 0 and i != len(sentence) - 1 and sentence[i-1].isdigit() and sentence[i+1].isdigit():
            word += letter
        #the case when . is a period preceded by a letter/number
        elif letter == '.' and i == 1 and (sentence[i-1].isdigit() or 'א' <= sentence[i-1] <= 'י'):
            word += letter
        elif letter == '.' and i > 1 and (sentence[i-1].isdigit() or 'א' <= sentence[i-1] <= 'י') and sentence[i-2] == ' ':
            word += letter
        elif letter in marks: #a mark that don't any of the above criteria will be a single token 
            if word:
                tokenized_sentence.append(word)
                word = ''
            if letter != ' ': 
                tokenized_sentence.append(letter)
        else:
            word += letter
    if word:
        tokenized_sentence.append(word)
    if len(tokenized_sentence) >= 4:
        tokenized_sentence = ' '.join(tokenized_sentence)
        return tokenized_sentence
    else:
        return None

#helper method for for extract_relevant_text that creates a sentence object and add it to the relevant protocol 
def sentence_handle(protocol, curr_speaker, paragraph_txt):
    #remove tags from sentence 
    open_an_brac = paragraph_txt.find('<')
    close_an_brac = paragraph_txt.find('>')
    if close_an_brac != len(paragraph_txt) - 1 and paragraph_txt[close_an_brac + 1] == '>':
        close_an_brac += 1 #in case it's >> not >
    while open_an_brac != -1 and close_an_brac != -1 and open_an_brac < close_an_brac: 
        paragraph_txt = paragraph_txt[:open_an_brac] + paragraph_txt[close_an_brac+1:]
        open_an_brac = paragraph_txt.find('<')
        close_an_brac = paragraph_txt.rfind('>')
    curr_sentence = '' #initializing empty sentence 
    for i,letter in enumerate(paragraph_txt): 
        if curr_sentence == '' and letter == ' ': #to avoid a sentence starting with spaces 
            continue
        curr_sentence += letter
        if letter in ['?', '!']: # ? ! always define end of a sentence 
            if sentence_validity(curr_sentence):
                curr_sentence = sentence_tokenize(curr_sentence)
                if curr_sentence:
                    sentence = Sentence(protocol.name, protocol.keneset, protocol.type, protocol.number, curr_speaker, curr_sentence)
                    protocol.add_sentence(curr_speaker, sentence)
            curr_sentence = ''
        elif letter == '.':
            #case 1: '.' is a decimal point, not considered as an end of a sentence
            if len(curr_sentence) > 0 and i != len(paragraph_txt) - 1 and curr_sentence[-1].isdigit() and paragraph_txt[i+1].isdigit():
               continue
            #case 2: '.' is a period, preceded by a number, not considered as an end of a sentence
            elif len(curr_sentence) > 2 and curr_sentence[-1].isdigit() and curr_sentence[-2] == ' ' and curr_sentence[-3] in [',' , ':']:
                continue
            elif len(curr_sentence) == 1 and curr_sentence[-1].isdigit():
                continue
            #case 3: '.' is a period, preceded by a letter, not considered as an end of a sentence
            elif len(curr_sentence) > 1 and  'א' <= curr_sentence[-1] <= 'י' and curr_sentence[-2] == ' ':
                continue
            #case 4: '.' is an end of a sentence
            else:
                if sentence_validity(curr_sentence):
                    curr_sentence = sentence_tokenize(curr_sentence)
                    if curr_sentence:
                        sentence = Sentence(protocol.name, protocol.keneset, protocol.type, protocol.number, curr_speaker, curr_sentence)
                        protocol.add_sentence(curr_speaker, sentence)
                curr_sentence = ''
    #handling the remaining text 
    if sentence_validity(curr_sentence): 
        curr_sentence = sentence_tokenize(curr_sentence)
        if curr_sentence:
            sentence = Sentence(protocol.name, protocol.keneset, protocol.type, protocol.number, curr_speaker, curr_sentence)
            protocol.add_sentence(curr_speaker, sentence)


#extracts the relevant sentences and arranges them according to protocol and speaker
def extract_relevant_text(file_content, protocol):
    #find the first and last paragraphs of the block where relevant texts 
    first_idx = find_starting_relevant(file_content)
    last_idx = find_last_relevant(file_content)
    curr_speaker = None
    for i,paragraph in enumerate(file_content.paragraphs[first_idx:last_idx]):
        #type 1 of sentences to exclude: those related to 'vote'. If it's a vote sentence, consider the text irrelevant until there's a speaker
        if 'הצבעה' in paragraph.text or 'ההצבעה' in paragraph.text or 'הצבעת' in paragraph.text:
            j = i+1
            while not file_content.paragraphs[first_idx+j].text:
                j += 1
            if 'בעד' in file_content.paragraphs[first_idx + j].text and 'נגד' in file_content.paragraphs[first_idx + j+1].text:
                curr_speaker = None
        #type 2 of sentences to exclude: those related to 'debate pause'. If it so, consider the text irrelevant until there's a speaker
        if 'הישיבה נפסקה' in paragraph.text:
            curr_speaker = None
        paragraph_txt = paragraph.text.strip()
        #meeting one of the first two if's making the paragraph as potential speaker's name
        if (paragraph_txt.endswith(':') or paragraph_txt.endswith(':>')) and is_underlined(paragraph):
            paragraph_txt_cleaned = speakerClean(paragraph_txt)[:-1].strip()
            first_string = re.search(r'\b\w+\b', paragraph_txt_cleaned)
            if len(paragraph_txt_cleaned.split()) < 6 and first_string and first_string.group(0) not in invalid_names:
                curr_speaker = paragraph_txt_cleaned
            elif curr_speaker: #deal with it as speech
                sentence_handle(protocol, curr_speaker, paragraph_txt)
        elif ':' in paragraph_txt:
            pot_curr_speaker = speakerClean(paragraph_txt)
            if pot_curr_speaker.endswith(':') and is_underlined(paragraph):
                paragraph_txt_cleaned = pot_curr_speaker[:-1].strip()
                first_string = re.search(r'\b\w+\b', paragraph_txt_cleaned)
                if len(paragraph_txt_cleaned.split()) < 6 and first_string and first_string.group(0) not in invalid_names:
                    curr_speaker = paragraph_txt_cleaned
                elif curr_speaker: #deal with it as speech
                    sentence_handle(protocol, curr_speaker, paragraph_txt)
            elif curr_speaker: #else deal with it as a speech 
                sentence_handle(protocol, curr_speaker, paragraph_txt)
        elif curr_speaker:
            sentence_handle(protocol, curr_speaker, paragraph_txt)

#Input: path to a folder
#Output: file names, paths, and contents for all .docx files in the folder
def read_files(folder):
    try:
        file_names = [file for file in os.listdir(folder) if (file[-5:] == '.docx' or file[-4:] == '.doc')]
        file_paths = {filename: os.path.join(folder, filename) for filename in file_names}
        file_contents = {}
        for file_name, file_path in file_paths.items():
            try: 
                file_contents[file_name] = Document(file_path)
            except Exception as e:
                print(f"Failed to read file {file_name} ")
                file_names.remove(file_name)
        return file_names, file_contents
    except FileNotFoundError:
        print(f"Folder {folder} is not exist")
    except Exception as e:
        print(f"Error: {e}")

#a function that creates the JSONL file 
#Input: list of protocols objects and file path where the JSONL file will be saved 
#Output: JSONL file creation, in which each line contain a sentence and its metadata 
def jsonl_make(protocols, file):
    with open(file, 'w', encoding = 'utf-8') as jsonl_file:
        for protocol in protocols:
            for speaker, sentences in protocol.sentences.items():
                for sentence in sentences:
                    sentence_data = {
                        "protocol_name": sentence.protocol_name,
                        "knesset_number": sentence.keneset,
                        "protocol_type": sentence.protocol_type,
                        "protocol_number": sentence.protocol_no,
                        "speaker_name": sentence.speaker,
                        "sentence_text": sentence.text
                    }
                    jsonl_file.write(json.dumps(sentence_data, ensure_ascii=False) + '\n')

def main():
    if len(sys.argv) != 3:
        print("Error: Incorrect # of arguments.\n")
        sys.exit(1)
    else:
        print("Creating the corpus ..\n")
    folder_path = sys.argv[1] 
    file = sys.argv[2] 
    file_names, file_contents = read_files(folder_path)
    protocols = []
    for file_name in file_names: 
        keneset_no, protocol_type = extract_metada_from_name(file_name)
        protocol_no = extract_metada_from_content(file_contents[file_name])
        protocol = Protocol(file_name, keneset_no, protocol_type, protocol_no)
        protocols.append(protocol)
        extract_relevant_text(file_contents[file_name], protocol)
    jsonl_make(protocols, file)

if __name__ == "__main__":
    main()